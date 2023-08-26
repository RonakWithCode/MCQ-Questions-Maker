import tkinter as tk
from tkinter import Label, Entry, Button
from docx import Document
from docx.shared import Pt, RGBColor


# Adinath public school
# Class 9th
# Marks: 80

def generate_layout():
    # Get user inputs
    mcq_questions = int(input("Enter the number of Questions "))

    # Create a Word document
    doc = Document()
    
    # Add school name, class, and marks at the top
    school_name = str(input("School "))
    class_infoinput = int(input("Class "))
    class_info = f"Class {class_infoinput}th"
    marks_infoinput = int(input("Marks: "))
    marks_info = f"Marks: {marks_infoinput}"
    
    title = doc.add_paragraph()
    title_run = title.add_run(f"{school_name}\n{class_info}\n{marks_info}")
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_heading('MCQ Type Question', level=0)

    for question in range(mcq_questions):
        doc.add_paragraph(f'Q{question + 1}: {input("question ")}')
        isliner = int(input("Type 1 or Some Else"))
        if (isliner == 1):
            doc.add_paragraph(f' a) {input("A  ")}  (   )     B) {input("B ")} (   )      C) {input("C ")} (   )      D) {input("D ")} (   )')
        else: 
            doc.add_paragraph(f' a) {input("A  ")}  (   )  \n B) {input("B ")} (   )  \n C) {input("C ")} (   )  \n D) {input("D ")} (   )')


    # Save the document
    doc.save('question_paper.docx')

generate_layout()